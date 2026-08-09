#!/usr/bin/env python3
"""Convert the 4 validated articles into individual Word documents."""
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ARTICLE_DIR = "/home/ubuntu/sunday_run_20260802/articles"
DOCX_DIR = "/home/ubuntu/sunday_run_20260802/docx"
os.makedirs(DOCX_DIR, exist_ok=True)

articles = [
    "20260802_article_1_funding.md",
    "20260802_article_2_exhibitions.md",
    "20260802_article_3_technology.md",
    "20260802_article_4_economics.md",
]

for article_file in articles:
    filepath = os.path.join(ARTICLE_DIR, article_file)
    with open(filepath, "r") as f:
        content = f.read()

    # Split into meta and body at Executive Summary
    parts = content.split("**Executive Summary**")
    meta_section = parts[0]
    body_and_notes = "**Executive Summary**" + parts[1] if len(parts) > 1 else ""

    # Remove Internal Evidence Note section for the Word doc
    body = body_and_notes.split("**Internal Evidence Note**")[0].strip()

    # Extract meta title
    title_match = re.search(r'Meta Title:\s*(.+)', meta_section)
    title = title_match.group(1).strip() if title_match else "Article"

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    doc.add_heading(title, level=1)

    # Process body
    lines = body.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('**Executive Summary**'):
            doc.add_heading('Executive Summary', level=2)
        elif line.startswith('This analysis synthesizes'):
            p = doc.add_paragraph(line)
            p.style = doc.styles['Normal']
            p.runs[0].italic = True
        else:
            # Remove bold markers for clean Word output
            clean_line = line.replace('**', '')
            doc.add_paragraph(clean_line)

    # Save
    docx_name = article_file.replace('.md', '.docx')
    doc.save(os.path.join(DOCX_DIR, docx_name))
    print(f"Saved: {docx_name}")

print("All articles converted to DOCX.")
