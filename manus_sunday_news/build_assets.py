#!/usr/bin/env python3
"""Build Sunday Tracker XLSX and Review Summary DOCX for 2 August 2026 weekly run."""
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from datetime import datetime

# --- Sunday Tracker ---
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Sunday Tracker"

header_font = Font(bold=True, size=11)
header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
header_font_white = Font(bold=True, size=11, color="FFFFFF")
thin_border = Border(
    left=Side(style='thin'), right=Side(style='thin'),
    top=Side(style='thin'), bottom=Side(style='thin')
)

headers = ["Article #", "Pillar", "Slug", "Word Count", "Compliance", "Hero Image", "S3 Upload", "CMS Post", "Status"]
for col, h in enumerate(headers, 1):
    cell = ws.cell(row=1, column=col, value=h)
    cell.font = header_font_white
    cell.fill = header_fill
    cell.alignment = Alignment(horizontal='center')
    cell.border = thin_border

articles = [
    ["1", "Funding & Capital Allocation", "august-2026-robotics-funding-infrastructure-capital", "1229", "ALL PASS", "20260802_hero_funding.jpg", "Pending", "Pending", "Ready"],
    ["2", "Global Exhibitions & Ecosystem", "august-2026-exhibitions-conventions-procurement-events", "1246", "ALL PASS", "20260802_hero_exhibitions.jpg", "Pending", "Pending", "Ready"],
    ["3", "Technology Developments", "august-2026-technology-hardware-reality-check", "1252", "ALL PASS", "20260802_hero_technology.jpg", "Pending", "Pending", "Ready"],
    ["4", "Pricing & Commercial Economics", "august-2026-robot-economics-regulation-sets-price", "1225", "ALL PASS", "20260802_hero_economics.jpg", "Pending", "Pending", "Ready"],
]

for row_idx, row_data in enumerate(articles, 2):
    for col_idx, val in enumerate(row_data, 1):
        cell = ws.cell(row=row_idx, column=col_idx, value=val)
        cell.border = thin_border
        cell.alignment = Alignment(horizontal='center')

# Adjust column widths
for col in ws.columns:
    max_length = max(len(str(cell.value or "")) for cell in col)
    ws.column_dimensions[col[0].column_letter].width = min(max_length + 4, 50)

wb.save("/home/ubuntu/sunday_run_20260802/assets/20260802_Sunday_Tracker.xlsx")
print("Sunday Tracker saved.")

# --- Review Summary ---
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

doc.add_heading('RobotAIGeek Sunday Wrap-Up: Review Summary', level=1)
doc.add_paragraph(f'Run Date: 2 August 2026 (Asia/Shanghai)')
doc.add_paragraph(f'Run Type: Standard Weekly Wrap-Up')
doc.add_paragraph(f'Coverage Period: 26 July - 1 August 2026')
doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M UTC")}')

doc.add_heading('Compliance Results', level=2)
doc.add_paragraph('All 4 articles passed the automated mechanical compliance validator on all 14 checks.')

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Article'
hdr[1].text = 'Word Count'
hdr[2].text = 'Compliance'
hdr[3].text = 'Status'
data = [
    ('1: Funding', '1229', 'ALL PASS', 'Ready'),
    ('2: Exhibitions', '1246', 'ALL PASS', 'Ready'),
    ('3: Technology', '1252', 'ALL PASS', 'Ready'),
    ('4: Economics', '1225', 'ALL PASS', 'Ready'),
]
for i, row_data in enumerate(data, 1):
    for j, val in enumerate(row_data):
        table.rows[i].cells[j].text = val

doc.add_heading('Deduplication Notes', level=2)
doc.add_paragraph(
    'The following topics from the reporting period were already covered on robotaigeek.com and were '
    'aggregated rather than repeated: Tesla VR treadmill (Jul 28 news), Korea 300kg gripper (Jul 28 news), '
    'China open-source reasoning model (Jul 28 news), Android creator robot OS (Jul 29 news), '
    'German gas-spring maker (Jul 29 news), Harper Adams farm robots (Jul 31 news), '
    'and the China IPO Gauntlet article (Aug 1). These were referenced for context only.'
)

doc.add_heading('Variance Engine Assignments', level=2)
ve_table = doc.add_table(rows=5, cols=4)
ve_table.style = 'Table Grid'
ve_hdr = ve_table.rows[0].cells
ve_hdr[0].text = 'Article'
ve_hdr[1].text = 'Opening Style'
ve_hdr[2].text = 'Persona'
ve_hdr[3].text = 'Closing Style'
ve_data = [
    ('1: Funding', 'Data Lead', 'Market Realist', 'Declarative Verdict'),
    ('2: Exhibitions', 'Event Anchor', 'Supply Chain Operator', 'Forward-Looking Signpost'),
    ('3: Technology', 'Structural Observation', 'Technical Translator', 'Scene Close'),
    ('4: Economics', 'Tension Frame', 'Commercial Operator', 'Statistic Close'),
]
for i, row_data in enumerate(ve_data, 1):
    for j, val in enumerate(row_data):
        ve_table.rows[i].cells[j].text = val

doc.add_heading('XLSX Assets', level=2)
doc.add_paragraph('No standalone XLSX assets required for this weekly run. The Sunday Tracker XLSX is the only spreadsheet deliverable.')

doc.add_heading('Notes', level=2)
doc.add_paragraph('No pre-staged companion article for this run date. The 26 July companion article was delivered in the prior run.')

doc.save("/home/ubuntu/sunday_run_20260802/assets/20260802_Review_Summary.docx")
print("Review Summary saved.")
