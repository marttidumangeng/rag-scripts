#!/usr/bin/env python3
"""Deterministic Word-first assembler for ARPI dedicated-topic articles.

Builds the v8-standard .docx package (cover page, metadata table, styled
body, embedded table PNGs, internal-only Evidence Source Notes closing
page, disclaimer) from structured inputs, then re-opens the result to
run the open-and-render test. Replaces per-run ad-hoc python-docx code.

Inputs:
  --assignment assignment.json   (from assignment_lookup.py)
  --body body.md                 publishable body: '## ' lines become H2
                                 headings; '{{table:file.png}}' on its own
                                 line embeds that PNG; '* ' lines become
                                 bullets; **bold** is honored; everything
                                 else is a paragraph. No inline citations.
  --evidence evidence.md         numbered Evidence Source Notes entries
                                 (internal-only; goes on the closing page)
  --meta meta.json               {"headline", "dateline", "hero_credit",
                                  "lane", "slug", "summary", "category",
                                  "tags", "meta_title", "meta_description"}
  --hero hero.jpg                16:9 JPG for the cover
  --out-dir .                    output directory

The output filename comes from assignment.docx_filename. Exit 0 only if
the built file re-opens and every required component is present.
Requires python-docx (pip install python-docx).
"""
import argparse
import datetime as dt
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

HEADING_COLOR = RGBColor(0x1F, 0x38, 0x64)  # dark blue
DISCLAIMER = (
    "This article is provided for general information purposes only and "
    "does not constitute investment, procurement, legal, or engineering "
    "advice. Figures reflect the sources listed above as of the stated "
    "information cut-off and may change without notice.")
TABLE_RE = re.compile(r"^\{\{table:(.+?)\}\}\s*$")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def add_runs(par, text):
    """Minimal **bold** support; everything else plain."""
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        par.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])


def style_heading(par, size=14):
    for run in par.runs:
        run.font.color.rgb = HEADING_COLOR
        run.font.size = Pt(size)
        run.font.bold = True


def build(assignment, meta, body_lines, evidence_text, hero, out_dir):
    doc = Document()
    table_images = [TABLE_RE.match(l.rstrip()).group(1)
                    for l in body_lines if TABLE_RE.match(l.rstrip())]

    # --- cover page ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(meta.get("headline") or assignment["title"])
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = HEADING_COLOR

    dateline = doc.add_paragraph()
    dateline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dateline.add_run(meta.get("dateline") or dt.date.fromisoformat(
        assignment["run_date"]).strftime("%B %d, %Y"))
    dr.font.size = Pt(11)

    if hero and Path(hero).exists():
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(hero), width=Inches(6.5))
        credit = doc.add_paragraph()
        credit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = credit.add_run(meta.get("hero_credit", "Image: generated for "
                                                    "this article"))
        cr.font.size = Pt(9)
        cr.font.italic = True
    doc.add_page_break()

    # --- metadata page ---
    h = doc.add_paragraph()
    h.add_run("Upload Metadata")
    style_heading(h, 16)
    # Fields required by SKILL.md Section 9 metadata block, plus the upload
    # form fields. Values come from assignment.json where possible so they
    # cannot drift from the calendar row.
    rows = [
        ("Destination lane", meta.get("lane", "Article")),
        ("Title", meta.get("headline") or assignment["title"]),
        ("URI slug", meta.get("slug", "")),
        ("Summary", meta.get("summary", "")),
        ("SEO meta title", meta.get("meta_title", "")),
        ("SEO meta description", meta.get("meta_description", "")),
        ("RobotAIGeek category", meta.get("category", "")),
        ("Tags", ", ".join(meta.get("tags", []))
            if isinstance(meta.get("tags"), list) else meta.get("tags", "")),
        ("Asset class", assignment["asset_class"]),
        ("Post slot", f"Post {assignment['post_slot']}"),
        ("Scheduled date", assignment["run_date"]),
        ("Author", meta.get("author", "Aeon (user ID 211)")),
        ("Status", "Draft - editorial review required"),
        ("Hero image", assignment["hero_filename"]),
        ("Table images", ", ".join(table_images) if table_images else "None"),
        ("Publish action", "Editorial review required - do not "
                           "auto-publish"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for key, val in rows:
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(key).bold = True
        cells[1].paragraphs[0].add_run(str(val))
    doc.add_page_break()

    # --- article body ---
    missing_tables = []
    for line in body_lines:
        line = line.rstrip()
        if not line:
            continue
        tm = TABLE_RE.match(line)
        if tm:
            png = Path(tm.group(1))
            if png.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(png), width=Inches(6.5))
            else:
                missing_tables.append(str(png))
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            add_runs(p, line[3:].strip())
            style_heading(p)
        elif line.startswith("# "):
            continue  # article title lives on the cover page
        elif line.startswith(("* ", "- ")):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:].strip())
        else:
            p = doc.add_paragraph()
            add_runs(p, line)
    doc.add_page_break()

    # --- internal-only closing page ---
    h = doc.add_paragraph()
    h.add_run("Evidence Source Notes (Internal Only – Do Not Publish)")
    style_heading(h, 14)
    for line in evidence_text.splitlines():
        if line.strip():
            p = doc.add_paragraph()
            add_runs(p, line.strip())
            for run in p.runs:
                run.font.size = Pt(9)
    doc.add_paragraph()
    d = doc.add_paragraph()
    dr = d.add_run(DISCLAIMER)
    dr.font.size = Pt(9)
    dr.font.italic = True

    out_path = Path(out_dir) / assignment["docx_filename"]
    doc.save(out_path)
    return out_path, missing_tables


def render_test(path):
    """Re-open and verify required components exist."""
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(c.text for t in doc.tables
                           for row in t.rows for c in row.cells)
    checks = {
        "cover headline": bool(doc.paragraphs and doc.paragraphs[0].text),
        "metadata table": "Publish action" in table_text,
        "body headings": any(p.text and p.runs and p.runs[0].font.bold
                             and p.runs[0].font.color.rgb == HEADING_COLOR
                             for p in doc.paragraphs[3:]),
        "evidence section": "Evidence Source Notes (Internal Only" in text,
        "disclaimer": "general information purposes" in text,
    }
    return checks


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--assignment", required=True)
    ap.add_argument("--body", required=True)
    ap.add_argument("--evidence", required=True)
    ap.add_argument("--meta", required=True)
    ap.add_argument("--hero")
    ap.add_argument("--out-dir", default=".")
    args = ap.parse_args()

    assignment = json.loads(Path(args.assignment).read_text(encoding="utf-8"))
    meta = json.loads(Path(args.meta).read_text(encoding="utf-8"))
    body_lines = Path(args.body).read_text(encoding="utf-8").splitlines()
    evidence = Path(args.evidence).read_text(encoding="utf-8")

    out_path, missing = build(assignment, meta, body_lines, evidence,
                              args.hero, args.out_dir)
    checks = render_test(out_path)

    print(f"wrote {out_path}")
    ok = all(checks.values()) and not missing
    for name, passed in checks.items():
        print(f"  {'ok  ' if passed else 'FAIL'} {name}")
    for t in missing:
        print(f"  FAIL table PNG not found: {t}")
    print(f"RENDER TEST: {'PASS' if ok else 'FAIL'}")
    sys.exit(0 if ok else 1)


if __name__ == "__main__":
    main()
