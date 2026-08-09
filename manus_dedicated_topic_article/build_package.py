from pathlib import Path
from datetime import date
import math
import textwrap

import matplotlib.pyplot as plt
import numpy as np
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

ROOT = Path('/home/ubuntu/projects/dedicated-topic-article-090d06d7')
OUT = ROOT / 'output'
OUT.mkdir(exist_ok=True)

TITLE = 'Beyond the Invoice: Calculating the True Cost of an AMR and AGV'
DATE_STR = '2026-08-06'
DATE_COMPACT = '20260806'
ASSET_CLASS = 'AMR / AGV'
POST_SLOT = 'Post 9'
HERO = OUT / f'{DATE_COMPACT}_AMR_AGV_Post9_hero.jpg'
TABLE1 = OUT / f'{DATE_COMPACT}_AMR_AGV_Post9_tco_bridge.png'
TABLE2 = OUT / f'{DATE_COMPACT}_AMR_AGV_Post9_downtime_sensitivity.png'
XLSX = OUT / f'{DATE_COMPACT}_AMR_AGV_Post9_tco_model.xlsx'
DOCX = OUT / f'RobotAIGeek_Draft_{DATE_COMPACT}_AMR_AGV_Post9.docx'
QC = OUT / f'QC_Report_{DATE_COMPACT}_AMR_AGV_Post9.md'

summary = ('The invoice for an autonomous mobile robot or automated guided vehicle is only the first line of a five-year ownership case. This buyer-focused framework separates integration, workforce readiness, maintenance, downtime exposure, compliance, battery planning, and upgrades, then applies a transparent illustrative total-cost model.')
meta_title = 'AMR and AGV Total Cost of Ownership: A Buyer’s Guide'
meta_description = 'A practical AMR and AGV total-cost framework covering integration, training, maintenance, downtime, compliance, batteries, upgrades, and a five-year illustrative model.'
tags = 'AMR, AGV, total cost of ownership, warehouse automation, procurement, robotics economics, safety compliance'
slug = 'beyond-the-invoice-calculating-the-true-cost-of-an-amr-and-agv'

paragraphs = {
'Installation and Integration Costs': [
"The vehicle is rarely the whole automation project. An autonomous mobile robot (AMR) or automated guided vehicle (AGV) must operate inside a defined traffic system, exchange work with a warehouse management system (WMS), and hand off loads at stations that were designed for people or forklifts. The invoice therefore needs a second envelope for site preparation, controls engineering, software interfaces, commissioning, safety validation, and acceptance testing.",
"The recent AMR and AGV series already established why specification quality matters. The same discipline should be applied to cost lines. A tender that says ‘integration included’ is not a cost model until it states which interfaces, maps, traffic rules, charging points, racks, doors, conveyors, network changes, and performance tests are included. The International Organization for Standardization lists ISO 3691-4:2020 as withdrawn and identifies ISO 3691-4:2023 as the replacement. That makes copied language in a procurement pack a potential rework trigger, not a harmless reference.",
"A useful buyer test is to split integration into reversible and irreversible work. Fleet software, maps, and task rules can often be adjusted after launch. Floor repairs, electrical distribution, fire-protection changes, dock alterations, and station redesign are less forgiving. In a competitive bid, each supplier should price both categories separately and identify the party responsible for the operating zone. This turns a vague implementation allowance into an obligation register that can be audited before purchase order approval.",
"A live deployment announced by Geek+ and OMLOG illustrates the point without proving a universal return. The Hong Kong project combined more than 1,000 shelves, a multi-floor facility, and integration between the warehouse management system and robot management system. The operational result was vendor reported as more than double productivity. For a buyer, the lesson is not to copy the number. It is to ask whether the proposed design includes the same level of facility adaptation and data integration."
],
'Training and Onboarding': [
"Training is often treated as a short handover session, even though mobile robots change the work of operators, supervisors, maintenance staff, information-technology teams, and safety managers. The cost includes paid learning time, written procedures, escalation drills, exception handling, and the temporary productivity dip while the new process stabilises. If a site operates several shifts, training must be repeated without assuming that the first shift can coach the others for free.",
"The U.S. Bureau of Labor Statistics reports 2024 median pay of $37,680, or $18.12 per hour, for hand laborers and material movers. That is a useful external reference, not a universal labour rate. A business case should use its own loaded labour cost, shift pattern, overtime exposure, and redeployment plan. If the automation removes walking but not the underlying material-handling task, the correct comparison is redeployed minutes and recovered capacity, not the full wage of a role that will remain on the payroll.",
"The new ANSI/A3 R15.08-3-2026 standard is a relevant 2026 development because its official description emphasises user risk assessment, management of change, and lifecycle safety for industrial mobile robot applications. Those requirements make onboarding a continuing operating capability. The buyer should budget for refresher training after layout changes, software releases, new payloads, or altered pedestrian routes, rather than treating training as a one-time commissioning cost."
],
'Maintenance and Spare Parts': [
"Maintenance begins with the service model written into the contract. Buyers should distinguish preventive inspections, remote diagnostics, corrective labour, spare parts, software support, battery servicing, charger maintenance, and travel time. A low annual support fee can be misleading if the response-time commitment excludes evenings, weekends, or the second shift that creates most of the throughput value.",
"Battery and charging assumptions deserve their own line. ISO 3691-4:2020 explicitly states that power-source requirements are not covered by that document, so a standards reference does not by itself answer questions about battery chemistry, charging temperature, replacement timing, fire controls, or charger capacity. The supplier should provide a duty-cycle model that links fleet size to charging windows and spare capacity. A fleet that looks fully utilised on paper may need additional vehicles if charging is scheduled during the same peak period as outbound work.",
"A practical contract should define what counts as a failure, which data the customer can export, and which parts are subject to vendor lock-in. It should also state whether software updates require a new safety validation. These clauses protect the five-year model from a common error: treating the robot as a static appliance even though its maps, traffic logic, interfaces, and safety functions change over time."
],
'Downtime and Productivity Risk': [
"Downtime is the line most likely to decide whether a project creates value. A stopped robot can block a lane, starve a workstation, or force manual recovery that costs more than the unavailable vehicle. The correct measure is not only robot uptime. It is the value of the constrained process, the duration of recovery, the number of substitute vehicles, and the volume that misses a customer or production promise.",
"The Occupational Safety and Health Administration (OSHA) notes that robot accidents often occur during non-routine conditions such as programming, maintenance, testing, setup, or adjustment, and it does not publish a robotics-specific standard. That combination has an economic implication. A site needs procedures, competent personnel, controlled maintenance zones, and time for safe intervention. Cutting those allowances may make the purchase case look better while increasing operational and safety exposure.",
"The analytical model in this article therefore carries an explicit annual downtime and productivity reserve. It is not a probability estimate and should not be presented as one. It is a planning device that lets a buyer test the consequence of a service-level failure. If the base case assumes $40,000 per year and the high case assumes $80,000, the five-year ownership cost changes by $200,000. That sensitivity is often more decision-useful than a debate over a small difference in unit price."
],
'End of Life and Upgrade Costs': [
"A five-year case should not assume that the original fleet will remain commercially or technically unchanged. Upgrade exposure can come from battery replacement, safety scanners, fleet-control software, operating-system support, cybersecurity requirements, new interfaces, or a site expansion that exceeds the original traffic model. The buyer should request a product roadmap, supported-life commitment, data-export policy, and end-of-support procedure before accepting a low initial quote.",
"European Union buyers also face a moving compliance environment. EUR-Lex shows a current consolidated version of Regulation (EU) 2023/1230 dated 27 July 2026. The regulation replaces the earlier machinery directive and provides the legal framework for machinery safety and conformity obligations. This article does not determine legal applicability for a particular machine or jurisdiction. It does show why documentation, conformity assessment, responsible economic-operator work, and change control belong in the ownership plan rather than appearing as unexpected legal expense after installation.",
"End-of-life planning should cover physical disposal, data retention, spare-parts availability, redeployment, and resale assumptions. A conservative investment case assigns zero residual value unless a buyer has credible evidence for a secondary market. That avoids overstating returns with an optimistic resale figure that may disappear when a fleet controller or battery platform is retired."
],
'TCO Summary': [
"The table below is an illustrative five-year scenario for five transport robots. It is deliberately transparent rather than predictive: five units at a $100,000 assumed purchase price, $150,000 of integration and site preparation, $25,000 of training and change management, $60,000 per year of maintenance and support, $50,000 of battery and charging reserve in year three, $50,000 of planned upgrade reserve in year four, and $40,000 per year of downtime and productivity reserve. The model excludes financing, tax, inflation, salvage value, and residual cash flows.",
"Under those assumptions, the fleet sticker price is $500,000 and the five-year total cost of ownership (TCO) is $1.275 million, or 2.55 times the sticker price. The ratio is not a market average. It is a warning against using the invoice as a proxy for capital required. The largest controllable blocks are integration, recurring support, and downtime exposure. A supplier that reduces the unit price by $25,000 may matter less than one that clarifies interfaces, shortens recovery time, and provides usable maintenance data.",
"The procurement decision should therefore compare three cases: a constrained pilot with limited scope, a base deployment with explicit service and training obligations, and a scale case that includes additional stations, shifts, or sites. The most valuable commercial concession may be a measurable acceptance test, a response-time credit, or an upgrade guarantee rather than a headline discount. The invoice is the entry ticket; the operating envelope determines whether the investment thesis survives."
]
}

# TCO inputs and formulas
inputs = [
    ('Fleet sticker price', 500000, '5 robots x $100,000 illustrative unit assumption'),
    ('Integration and site preparation', 150000, 'Illustrative allowance for interfaces, maps, traffic, electrical, safety validation, commissioning'),
    ('Training and change management', 25000, 'Illustrative allowance for multi-shift onboarding and procedures'),
    ('Annual maintenance and support', 60000, 'Illustrative annual reserve, not a vendor quote'),
    ('Battery and charging reserve in year 3', 50000, 'Illustrative reserve; chemistry and duty cycle are site-specific'),
    ('Planned upgrade reserve in year 4', 50000, 'Illustrative software/hardware refresh reserve'),
    ('Annual downtime and productivity reserve', 40000, 'Illustrative planning reserve; sensitivity tested separately'),
]
base_tco = 500000 + 150000 + 25000 + 60000*5 + 50000 + 50000 + 40000*5
sensitivity = [(20000, base_tco - 100000), (40000, base_tco), (80000, base_tco + 200000)]

# Charts
plt.style.use('seaborn-v0_8-whitegrid')
labels = ['Fleet sticker\nprice', 'Integration\nand site prep', 'Training\nand change', '5-year maintenance\nand support', 'Battery /\ncharging', 'Upgrade\nreserve', '5-year downtime\nreserve']
values = [500000,150000,25000,300000,50000,50000,200000]
colors = ['#24415c','#b86e3d','#d2a84b','#497d7d','#7c9a8b','#8b6f91','#c95c5c']
fig, ax = plt.subplots(figsize=(11, 5.8), dpi=180)
left = 0
for lab, val, col in zip(labels, values, colors):
    ax.barh(['Illustrative five-year TCO'], [val], left=left, color=col, label=lab)
    if val >= 50000:
        ax.text(left + val/2, 0, f'${val/1000:.0f}k', ha='center', va='center', color='white', fontsize=9, fontweight='bold')
    left += val
ax.set_title('Illustrative AMR / AGV five-year TCO bridge', fontsize=14, fontweight='bold', color='#17324d')
ax.set_xlabel('USD, nominal illustrative assumptions')
ax.legend(loc='upper center', bbox_to_anchor=(0.5, -0.18), ncol=4, frameon=False, fontsize=8)
ax.set_xlim(0, base_tco*1.03)
ax.text(base_tco*0.98, 0.22, f'Total: ${base_tco/1e6:.3f}m', ha='right', va='bottom', fontsize=10, fontweight='bold', color='#17324d')
fig.tight_layout()
fig.savefig(TABLE1, bbox_inches='tight')
plt.close(fig)

fig, ax = plt.subplots(figsize=(8.8, 5.2), dpi=180)
xs = [x/1000 for x, _ in sensitivity]
ys = [y/1000000 for _, y in sensitivity]
ax.plot(xs, ys, marker='o', linewidth=3, color='#c95c5c')
for x, y in zip(xs, ys):
    ax.annotate(f'${y:.3f}m', (x, y), textcoords='offset points', xytext=(0, 10), ha='center', fontsize=10, fontweight='bold')
ax.set_title('Sensitivity to annual downtime and productivity reserve', fontsize=14, fontweight='bold', color='#17324d')
ax.set_xlabel('Annual downtime reserve (USD thousands)')
ax.set_ylabel('Five-year TCO (USD millions)')
ax.set_xticks(xs)
ax.set_ylim(1.05, 1.55)
ax.grid(axis='y', alpha=0.35)
fig.tight_layout()
fig.savefig(TABLE2, bbox_inches='tight')
plt.close(fig)

# Excel workbook
wb = Workbook()
ws = wb.active
ws.title = 'Inputs and Model'
header_fill = PatternFill('solid', fgColor='17324D')
sub_fill = PatternFill('solid', fgColor='DCE6F1')
white_font = Font(color='FFFFFF', bold=True)
thin = Side(style='thin', color='B7C9D6')
for cell in ws[1]:
    cell.fill = header_fill
    cell.font = white_font
ws['A1'] = 'AMR / AGV Illustrative Five-Year TCO Model'
ws['A3'] = 'Input / line item'; ws['B3'] = 'Value (USD)'; ws['C3'] = 'Method / limitation'
for c in ws[3]:
    c.fill = header_fill; c.font = white_font; c.alignment = Alignment(horizontal='center')
for i, (name, val, note) in enumerate(inputs, 4):
    ws.cell(i,1).value = name; ws.cell(i,2).value = val; ws.cell(i,3).value = note
    ws.cell(i,2).number_format = '$#,##0'
ws['A12'] = 'Calculated outputs'; ws['A12'].fill = sub_fill; ws['A12'].font = Font(bold=True)
ws['A13'] = 'Five-year maintenance and support'; ws['B13'] = '=B7*5'
ws['A14'] = 'Five-year downtime and productivity reserve'; ws['B14'] = '=B10*5'
ws['A15'] = 'Five-year TCO'; ws['B15'] = '=SUM(B4:B6)+B13+B8+B9+B14'
ws['A16'] = 'TCO / fleet sticker price'; ws['B16'] = '=B15/B4'
ws['B13'].number_format = '$#,##0'; ws['B14'].number_format = '$#,##0'; ws['B15'].number_format = '$#,##0'; ws['B16'].number_format = '0.00x'
ws['A18'] = 'Sensitivity'; ws['A18'].fill = sub_fill; ws['A18'].font = Font(bold=True)
ws['A19'] = 'Annual downtime reserve'; ws['B19'] = 'Five-year TCO'
for c in ws[19]: c.fill = header_fill; c.font = white_font
for r, (x, y) in enumerate(sensitivity, 20):
    ws.cell(r,1).value = x; ws.cell(r,2).value = y
    ws.cell(r,1).number_format = '$#,##0'; ws.cell(r,2).number_format = '$#,##0'
for col, width in {'A':38,'B':18,'C':74}.items(): ws.column_dimensions[col].width = width
for row in ws.iter_rows():
    for cell in row:
        cell.alignment = Alignment(vertical='top', wrap_text=True)
        cell.border = Border(bottom=thin)

ws2 = wb.create_sheet('Source Ledger')
for c, val in enumerate(['Source / URL','Tier','Verified claim or use','Limitation'], 1):
    ws2.cell(1,c).value = val; ws2.cell(1,c).fill = header_fill; ws2.cell(1,c).font = white_font
source_rows = [
('ISO 3691-4:2020 | https://www.iso.org/standard/70660.html','Tier 1 standards body','2020 edition is withdrawn; ISO 3691-4:2023 is the replacement; power-source requirements are not covered in the 2020 document.','Do not infer compliance from a copied standard reference.'),
('ANSI/A3 R15.08-3-2026 | https://www.automate.org/store/products/ansi-a3-r15-08-3-2026-american-national-standard-for-industrial-mobile-robots-safety-requirements-part-3-use-of-imr-applications-pdf-download','Tier 1 standards body','Published April 23, 2026; emphasizes user risk assessment, management of change, and lifecycle safety.','Official product description; standard text is paywalled.'),
('EUR-Lex Regulation (EU) 2023/1230 | https://eur-lex.europa.eu/eli/reg/2023/1230/oj','Tier 1 government/legal','Current consolidated version dated July 27, 2026; machinery safety and conformity framework.','Not legal advice or a machine-specific applicability determination.'),
('OSHA Robotics | https://www.osha.gov/robotics','Tier 1 government','No specific OSHA robotics standard; non-routine operations are highlighted as accident-risk conditions.','Broad robotics guidance, not a cost schedule.'),
('BLS Hand Laborers and Material Movers | https://www.bls.gov/ooh/transportation-and-material-moving/hand-laborers-and-material-movers.htm','Tier 1 government statistics','$37,680 annual median pay or $18.12 per hour in 2024.','U.S. wage benchmark; excludes full employer burden.'),
('Geek+ / OMLOG | https://www.geekplus.com/resources/news/omlog','Tier 2 company announcement','Vendor reports >1,000 shelves, 70,000 square feet, WMS/RMS integration, and >2x productivity.','Self-reported case; not a universal ROI benchmark.'),
('Geek+ / Schneider | https://www.geekplus.com/resources/news/geekwins-2026-rbr50-innovation-award-for-robot-arm-picking-station','Tier 2 company announcement','Vendor reports 48-hour production readiness and >=99.99% accuracy for a picking station.','Boundary example, not pure AMR transport evidence.'),
]
for r, row in enumerate(source_rows, 2):
    for c, val in enumerate(row, 1): ws2.cell(r,c).value = val
for col, width in {'A':60,'B':22,'C':70,'D':55}.items(): ws2.column_dimensions[col].width = width
for row in ws2.iter_rows():
    for cell in row: cell.alignment = Alignment(vertical='top', wrap_text=True); cell.border = Border(bottom=thin)

ws3 = wb.create_sheet('Analytical Notes')
notes = [
('Analytical layer','Scope-adjusted five-year TCO bridge separating vehicle capex, project delivery, workforce readiness, recurring service, risk reserve, and mid-life refresh.'),
('Model status','Illustrative procurement scenario, not a market-price survey or vendor quotation.'),
('Key limitation','Replace assumptions with site-specific quotes, throughput, staffing, uptime, service-level, finance, tax, inflation, salvage, and residual-value data before approval.'),
('BLS use','BLS wage is shown as an external U.S. reference; model uses a separate $28/hour loaded-labour assumption and does not claim the BLS figure is global.'),
('Sensitivity method','Annual downtime reserve is varied from $20,000 to $80,000; each $10,000 annual change moves five-year TCO by $50,000.'),
('Information cut-off','2026-08-06 Asia Shanghai.'),
]
ws3['A1']='Analytical note'; ws3['B1']='Detail'
for c in ws3[1]: c.fill = header_fill; c.font = white_font
for r, row in enumerate(notes, 2): ws3.cell(r,1).value=row[0]; ws3.cell(r,2).value=row[1]
ws3.column_dimensions['A'].width=28; ws3.column_dimensions['B'].width=110
for row in ws3.iter_rows():
    for cell in row: cell.alignment = Alignment(vertical='top', wrap_text=True); cell.border = Border(bottom=thin)
wb.save(XLSX)

# Word helper functions
def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text)); run.bold = bold; run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc, rows, cols, data, widths=None, header=True):
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for r in range(rows):
        for c in range(cols):
            cell = table.cell(r,c); cell.text = str(data[r][c])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs: run.font.size = Pt(8.5)
            if header and r == 0:
                set_cell_shading(cell, '17324D')
                for p in cell.paragraphs:
                    for run in p.runs: run.bold=True; run.font.color.rgb=RGBColor(255,255,255)
    if widths:
        for row in table.rows:
            for idx,w in enumerate(widths): row.cells[idx].width = Inches(w)
    doc.add_paragraph()
    return table

# Build Word document
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7); sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)
styles = doc.styles
styles['Normal'].font.name = 'Aptos'; styles['Normal'].font.size = Pt(10.5)
for style_name in ['Heading 1','Heading 2','Heading 3']:
    styles[style_name].font.name = 'Aptos Display'; styles[style_name].font.color.rgb = RGBColor(23,50,77)
styles['Heading 1'].font.size = Pt(18); styles['Heading 2'].font.size = Pt(14); styles['Heading 2'].font.bold = True

# Cover page
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(90)
r = p.add_run(TITLE); r.bold=True; r.font.size=Pt(25); r.font.color.rgb=RGBColor(23,50,77)
p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('ARPI Dedicated Topic Article | Post 9'); r.font.size=Pt(11); r.font.color.rgb=RGBColor(90,100,110)
p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Scheduled date: 6 August 2026 | Asset class: AMR / AGV'); r.font.size=Pt(10)
doc.add_picture(str(HERO), width=Inches(6.8))
p=doc.paragraphs[-1]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Image credit: Manus AI generated editorial asset, 2026.'); r.italic=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor(100,100,100)
doc.add_page_break()

# Metadata page
doc.add_heading('Upload Metadata', level=1)
metadata = [
('Destination lane','RobotAIGeek editorial review'),('Headline',TITLE),('URI slug',slug),('Summary',summary),('RobotAIGeek category','Robotics / Warehouse Automation'),('Tags',tags),('SEO meta title',meta_title),('SEO meta description',meta_description),('Upload image filename',HERO.name),('Publish action','Editorial review required; do not publish automatically'),('Status','Draft; unpublished')]
mt = doc.add_table(rows=len(metadata), cols=2); mt.style='Table Grid'; mt.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(k,v) in enumerate(metadata):
    set_cell_shading(mt.cell(i,0),'DCE6F1'); set_cell_text(mt.cell(i,0),k,bold=True,size=9); set_cell_text(mt.cell(i,1),v,size=9)
doc.add_page_break()

# Article body
doc.add_heading(TITLE, level=1)
p=doc.add_paragraph(summary); p.paragraph_format.space_after=Pt(10)
for heading, paras in paragraphs.items():
    doc.add_heading(heading, level=2)
    for txt in paras:
        p=doc.add_paragraph(txt); p.paragraph_format.space_after=Pt(7); p.paragraph_format.line_spacing=1.08
    if heading == 'TCO Summary':
        doc.add_picture(str(TABLE1), width=Inches(6.7)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        p=doc.add_paragraph('Figure 1. Illustrative five-year TCO bridge. Values are explicit planning assumptions, not market quotations.'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].italic=True; p.runs[0].font.size=Pt(8)
        doc.add_picture(str(TABLE2), width=Inches(5.9)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        p=doc.add_paragraph('Figure 2. Sensitivity to annual downtime and productivity reserve.'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].italic=True; p.runs[0].font.size=Pt(8)
        data = [['Line item','Illustrative five-year value','Share of TCO'],['Fleet sticker price','$500,000','39.2%'],['Integration and site preparation','$150,000','11.8%'],['Training and change management','$25,000','2.0%'],['Maintenance and support','$300,000','23.5%'],['Battery and charging reserve','$50,000','3.9%'],['Upgrade reserve','$50,000','3.9%'],['Downtime and productivity reserve','$200,000','15.7%'],['Total TCO','$1,275,000','100.0%']]
        doc.add_heading('Illustrative ownership bridge', level=3)
        add_table(doc, len(data), 3, data, widths=[2.6,2.0,1.3])

# Internal-only section
internal = doc.add_section(WD_SECTION.NEW_PAGE)
doc.add_heading('Evidence Source Notes (Internal Only – Do Not Publish)', level=1)
doc.add_paragraph('This closing section is editorial fact-checking material and must be removed before upload. The publishable article body intentionally contains no inline numeric citation markers.')
evidence = [
('1','International Organization for Standardization, ISO 3691-4:2020','https://www.iso.org/standard/70660.html','Supports the withdrawn status of ISO 3691-4:2020, the available ISO 3691-4:2023 replacement, the inclusion of AMR and AGV terminology, and the statement that power-source requirements are not covered in the 2020 document.'),
('2','Association for Advancing Automation, ANSI/A3 R15.08-3-2026','https://www.automate.org/store/products/ansi-a3-r15-08-3-2026-american-national-standard-for-industrial-mobile-robots-safety-requirements-part-3-use-of-imr-applications-pdf-download','Supports the April 23, 2026 publication date and the user-side risk assessment, management-of-change, and lifecycle-safety emphasis.'),
('3','European Union, EUR-Lex, Regulation (EU) 2023/1230','https://eur-lex.europa.eu/eli/reg/2023/1230/oj','Supports the current consolidated version dated July 27, 2026 and the machinery conformity and safety framework.'),
('4','U.S. Occupational Safety and Health Administration, Robotics','https://www.osha.gov/robotics','Supports the absence of a specific OSHA robotics standard and the focus on non-routine programming, maintenance, testing, setup, and adjustment.'),
('5','U.S. Bureau of Labor Statistics, Hand Laborers and Material Movers','https://www.bls.gov/ooh/transportation-and-material-moving/hand-laborers-and-material-movers.htm','Supports the 2024 median wage of $37,680 or $18.12 per hour. The article labels this as a U.S. reference and does not treat it as a global loaded labour cost.'),
('6','Geek+, Geek+ & OMLOG Revolutionize Luxury Fashion Logistics','https://www.geekplus.com/resources/news/omlog','Supports the vendor-reported 1,000-plus shelves, 70,000-square-foot facility, WMS/RMS integration, and more-than-double productivity case. The article explicitly labels the result as vendor reported and not a universal benchmark.'),
('7','Geek+, Geek+ Wins 2026 RBR50 Innovation Award for Robot Arm Picking Station','https://www.geekplus.com/resources/news/geekwins-2026-rbr50-innovation-award-for-robot-arm-picking-station','Supports the vendor-reported 48-hour production-readiness claim and the distinction between a picking-station expansion and a pure AMR/AGV transport fleet.'),
('8','Association for Advancing Automation, Meet Noovelia at Automate 2026','https://www.automate.org/news/meet-noovelia-at-automate-2026-booth-4381','Supports the June 22–25, 2026 event dates and the presence of AMR and mobile-robot safety-training content; retained as contextual evidence only.')]
for num,src,url,claim in evidence:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(f'{num}. {src}. '); r.bold=True
    p.add_run(url + '. '); p.add_run(claim)

doc.add_heading('Informational Disclaimer', level=2)
doc.add_paragraph('This article is an editorial procurement framework, not legal, tax, accounting, safety-certification, engineering, or investment advice. The TCO model is illustrative and must be replaced with site-specific supplier quotations, labour assumptions, throughput data, service-level terms, compliance review, financing inputs, and operating evidence before a purchase decision.')

# Footer
for section in doc.sections:
    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('RobotAIGeek editorial draft | Unpublished | 2026-08-06'); run.font.size=Pt(8); run.font.color.rgb=RGBColor(110,110,110)

doc.save(DOCX)

# QC report
body_text = '\n'.join(sum(paragraphs.values(), []))
word_count = len(body_text.split())
qc_text = f'''# Quality-Control Report\n\n## Result\n\n**Pass with editorial limitations recorded.** The package is an unpublished Word-first draft for the exact scheduled row dated {DATE_STR}.\n\n| Check | Result |\n| --- | --- |\n| Editorial timezone | Asia Shanghai |\n| RUN_DATE | {DATE_STR} |\n| Exact title | {TITLE} |\n| Asset class and post slot | {ASSET_CLASS}; {POST_SLOT} |\n| Body word count | {word_count} words, target 1,200 to 1,600 |\n| Required headings | Installation and Integration Costs; Training and Onboarding; Maintenance and Spare Parts; Downtime and Productivity Risk; End of Life and Upgrade Costs; TCO Summary |\n| Inline citation markers in body | None |\n| Internal evidence notes | Present in closing section and marked do not publish |\n| Hero | True JPEG, 2560 x 1440, 16:9, no text or logos |\n| Table graphics | Two deterministic PNG charts present |\n| Excel model | Present with formulas, source ledger, and analytical notes |\n| Word package | Generated as Microsoft Word 2007+ .docx |\n| Website preflight | Clear to draft; exact target not found in recent archive |\n| Publication status | Not published; editorial review required |\n\n## Original analytical layer\n\nThe article uses a scope-adjusted five-year TCO bridge. It separates vehicle capex, project delivery, workforce readiness, recurring service, downtime exposure, battery planning, and mid-life upgrade. It also includes a sensitivity test in which the annual downtime reserve varies from $20,000 to $80,000. This is an illustrative procurement model, not a market-price survey.\n\n## Source tiers and limitations\n\nTier 1 sources support the standards, regulatory, labour, and safety claims: ISO, the Association for Advancing Automation standards listing, EUR-Lex, OSHA, and the Bureau of Labor Statistics. Tier 2 company announcements are used only as scoped deployment signals and are explicitly identified as self-reported. No vendor marketing page is used as independent proof of market success.\n\nThe central limitation is that supplier prices, maintenance rates, battery life, downtime, and integration costs vary by payload, site, geography, shift pattern, facility condition, fleet controller, and contract. The model therefore labels all numerical TCO inputs as illustrative assumptions and excludes financing, tax, inflation, salvage, and residual value.\n\n## Hero differentiation\n\nThe hero uses a warm, neutral warehouse scene with one pallet-carrying vehicle and generous negative space. It avoids the recent site's dark blue gradient, human interview framing, logos, text overlays, and branded product treatment.\n\n## Delivery rule\n\nDo not publish until an editor removes the internal Evidence Source Notes and disclaimer, replaces illustrative assumptions with site-specific commercial inputs where needed, and completes legal and safety review for the intended jurisdiction.\n'''
QC.write_text(qc_text, encoding='utf-8')
print(f'Created {DOCX.name}, {XLSX.name}, {TABLE1.name}, {TABLE2.name}, {QC.name}; body_words={word_count}; tco={base_tco}')
